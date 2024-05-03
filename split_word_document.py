from docx import Document

def count_words(doc):
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    return word_count

def split_document(doc, words_per_file):
    new_doc = Document()
    words_current = 0
    for para in doc.paragraphs:
        words_in_para = len(para.text.split())
        if words_current + words_in_para > words_per_file:
            new_doc.save(f'split_{words_per_file}.docx')
            new_doc = Document()
            words_current = 0
        new_doc.add_paragraph(para.text)
        words_current += words_in_para
    new_doc.save(f'split_{words_per_file}.docx')

# Load your document
doc = Document('path_to_your_document.docx')
print(f'Total word count: {count_words(doc)}')

# Split document every 500 words
split_document(doc, 500)

